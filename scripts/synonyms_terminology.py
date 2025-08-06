import re
from docx import Document
from pathlib import Path

# ===== Paths =====
BASE_DIR = Path(__file__).resolve().parent.parent
INPUT_FILE = BASE_DIR / "input_data" / "ICT_toppics.docx"
OUTPUT_FILE = BASE_DIR / "process_data" / "ICT_toppics_expresions_edit.docx"

OUTPUT_FILE.parent.mkdir(exist_ok=True)

# ===== Config =====
PREFER_ABBREVIATION = True  # True = keep acronym, False = keep full term

# ===== Manual ICT-specific synonyms =====
manual_synonyms = {
    "Artificial Intelligence": "AI",
    "AI": "AI",
    "Machine Learning": "ML",
    "ML": "ML",
    "Natural Language Processing": "NLP",
    "NLP": "NLP",
    "Customer Satisfaction": "CSAT",
    "CSAT": "CSAT",
    "Customer Effort Score": "CES",
    "CES": "CES",
    "Net Promoter Score": "NPS",
    "NPS": "NPS",
    "Track NPS": "NPS",
    "Internet of Things": "IoT",
    "IoT": "IoT",
    "Structured Query Language": "SQL",
    "SQL": "SQL",
    "us": "USA",
    "usa": "USA",
    "the california consumer privacy act": "CPRA",
    "cpra": "CPRA",
    "the european data protection board": "EDPB",
    "edpb": "EDPB",
}

def build_synonym_map(doc):
    """Extract acronym–full form pairs from the document."""
    synonym_map = dict(manual_synonyms)  # Start with manual entries
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Full form (ACRONYM)
        match1 = re.findall(r"\b([A-Z][A-Za-z ]{2,})\s*\(([A-Z]{2,})\)", text)
        for full_form, acronym in match1:
            canonical = acronym if PREFER_ABBREVIATION else full_form
            synonym_map[full_form.lower()] = canonical
            synonym_map[acronym.lower()] = canonical

        # ACRONYM (Full form)
        match2 = re.findall(r"\b([A-Z]{2,})\s*\(([A-Z][A-Za-z ]{2,})\)", text)
        for acronym, full_form in match2:
            canonical = acronym if PREFER_ABBREVIATION else full_form
            synonym_map[full_form.lower()] = canonical
            synonym_map[acronym.lower()] = canonical

    return synonym_map

def replace_synonyms_in_doc(doc, synonym_map):
    """Replace synonyms in all paragraphs."""
    for para in doc.paragraphs:
        for term, canonical in synonym_map.items():
            para.text = re.sub(rf"\b{re.escape(term)}\b", canonical, para.text, flags=re.IGNORECASE)
    return doc

def get_synonym_map():
    """Return the synonym map from the INPUT_FILE."""
    if not INPUT_FILE.exists():
        print(f"❌ Input file not found: {INPUT_FILE}")
        return {}
    doc = Document(INPUT_FILE)
    return build_synonym_map(doc)

# ===== Create global synonym_map on import =====
synonym_map = get_synonym_map()

if __name__ == "__main__":
    if not INPUT_FILE.exists():
        print(f"❌ Input file not found: {INPUT_FILE}")
        exit(1)

    print("📌 Synonym map found:", synonym_map)

    # Replace terms in document
    doc = Document(INPUT_FILE)
    cleaned_doc = replace_synonyms_in_doc(doc, synonym_map)

    # Save cleaned file
    cleaned_doc.save(OUTPUT_FILE)
    print(f"✅ Cleaned document saved to: {OUTPUT_FILE}")
